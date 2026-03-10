import logging
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
import hashlib
from datetime import datetime
import json

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import xml.etree.ElementTree as ET
from bs4 import BeautifulSoup
import pytesseract
from PIL import Image

# ML libraries
from sentence_transformers import SentenceTransformer

from ..models.document import Document, DocumentType, AccessLevel, DocumentMetadata, Entity, EntityType
from ..services.entity_extractor import EntityExtractor
from ..services.vector_store import VectorStore
from config.settings import settings


class DocumentProcessor:
    """Service for processing and indexing documents."""
    
    def __init__(self):
        """Initialize document processor."""
        self.logger = logging.getLogger(__name__)
        self.embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL)
        self.entity_extractor = EntityExtractor()
        self.vector_store = VectorStore()
        
    async def process_document(self, file_path: str, access_level: AccessLevel = AccessLevel.PUBLIC) -> Optional[Document]:
        """
        Process a document and add it to the search index.
        
        Args:
            file_path: Path to the document file
            access_level: Access level for the document
            
        Returns:
            Processed Document object or None if processing failed
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                self.logger.error(f"File not found: {file_path}")
                return None
            
            # Determine document type
            doc_type = self._get_document_type(file_path)
            if not doc_type:
                self.logger.error(f"Unsupported file type: {file_path.suffix}")
                return None
            
            # Extract text content
            content = self._extract_text(file_path, doc_type)
            if not content:
                self.logger.error(f"Failed to extract text from: {file_path}")
                return None
            
            # Generate document ID
            doc_id = self._generate_document_id(file_path, content)
            
            # Extract metadata
            metadata = self._extract_metadata(file_path, doc_type, content)
            
            # Extract entities
            entities = self.entity_extractor.extract_entities(content)
            
            # Generate embedding
            embedding = self.embedding_model.encode(content).tolist()
            
            # Create document object
            document = Document(
                id=doc_id,
                filename=file_path.name,
                file_path=str(file_path),
                document_type=doc_type,
                access_level=access_level,
                content=content,
                metadata=metadata,
                entities=entities,
                embedding=embedding
            )
            
            # Add to vector store
            await self._add_to_vector_store(document)
            
            # Save processed document metadata
            await self._save_document_metadata(document)
            
            self.logger.info(f"Successfully processed document: {file_path.name}")
            return document
            
        except Exception as e:
            self.logger.error(f"Error processing document {file_path}: {str(e)}")
            return None
    
    def _get_document_type(self, file_path: Path) -> Optional[DocumentType]:
        """Determine document type from file extension."""
        extension = file_path.suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TEXT,
            '.docx': DocumentType.DOCX,
            '.xml': DocumentType.XML,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.png': DocumentType.IMAGE,
            '.gif': DocumentType.IMAGE,
            '.mp3': DocumentType.AUDIO,
            '.wav': DocumentType.AUDIO,
            '.mp4': DocumentType.VIDEO,
            '.avi': DocumentType.VIDEO
        }
        
        return type_mapping.get(extension)
    
    def _extract_text(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extract text content from document based on type."""
        try:
            if doc_type == DocumentType.PDF:
                return self._extract_text_from_pdf(file_path)
            elif doc_type == DocumentType.TEXT:
                return self._extract_text_from_txt(file_path)
            elif doc_type == DocumentType.DOCX:
                return self._extract_text_from_docx(file_path)
            elif doc_type == DocumentType.XML:
                return self._extract_text_from_xml(file_path)
            elif doc_type == DocumentType.IMAGE:
                return self._extract_text_from_image(file_path)
            else:
                self.logger.warning(f"Text extraction not implemented for type: {doc_type}")
                return ""
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            self.logger.error(f"Error reading PDF {file_path}: {str(e)}")
        
        return text.strip()
    
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            self.logger.error(f"Could not decode text file: {file_path}")
            return ""
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(str(file_path))
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            self.logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    def _extract_text_from_xml(self, file_path: Path) -> str:
        """Extract text from XML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            soup = BeautifulSoup(content, 'lxml')
            return soup.get_text(separator=' ', strip=True)
            
        except Exception as e:
            self.logger.error(f"Error reading XML {file_path}: {str(e)}")
            return ""

    def _extract_text_from_image(self, file_path: Path) -> str:
        """Extract text from Image file using OCR."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error performing OCR on image {file_path}: {str(e)}")
            return ""
    
    def _extract_metadata(self, file_path: Path, doc_type: DocumentType, content: str) -> DocumentMetadata:
        """Extract metadata from document."""
        try:
            stat = file_path.stat()
            
            # Basic metadata
            metadata = DocumentMetadata(
                title=file_path.stem,  # Filename without extension
                creation_date=datetime.fromtimestamp(stat.st_ctime),
                language="en",  # Default to English for POC
                file_size=stat.st_size,
                keywords=self._extract_keywords(content)
            )
            
            # Document type specific metadata
            if doc_type == DocumentType.PDF:
                metadata.page_count = self._get_pdf_page_count(file_path)
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error extracting metadata from {file_path}: {str(e)}")
            return DocumentMetadata()
    
    def _get_pdf_page_count(self, file_path: Path) -> Optional[int]:
        """Get page count for PDF files."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            self.logger.error(f"Error getting PDF page count: {str(e)}")
            return None
    
    def _extract_keywords(self, content: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from document content."""
        words = content.lower().split()
        
        # Filter out common stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
            'of', 'with', 'by', 'from', 'up', 'about', 'into', 'through', 'during',
            'before', 'after', 'above', 'below', 'between', 'among', 'is', 'are',
            'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does',
            'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can'
        }
        
        # Count word frequencies
        word_freq = {}
        for word in words:
            word = word.strip('.,!?;:"()[]{}')
            if len(word) > 3 and word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        return [word for word, freq in sorted_words[:max_keywords]]
    
    def _generate_document_id(self, file_path: Path, content: str) -> str:
        """Generate unique document ID."""
        # Create hash from file path and content
        hasher = hashlib.md5()
        hasher.update(str(file_path).encode('utf-8'))
        hasher.update(content[:1000].encode('utf-8'))  # First 1000 chars
        return hasher.hexdigest()
    
    async def _add_to_vector_store(self, document: Document):
        """Add document to vector store."""
        try:
            # Prepare metadata for vector store
            metadata = {
                'title': document.metadata.title or document.filename,
                'filename': document.filename,
                'document_type': document.document_type.value,
                'access_level': document.access_level.value,
                'creation_date': document.metadata.creation_date.isoformat() if document.metadata.creation_date else None,
                'language': document.metadata.language,
                'file_size': document.metadata.file_size,
                'page_count': document.metadata.page_count,
                'keywords': json.dumps(document.metadata.keywords),
                'entities': json.dumps([{
                    'text': e.text,
                    'entity_type': e.entity_type.value,
                    'confidence': e.confidence
                } for e in document.entities])
            }
            
            await self.vector_store.add_document(
                document_id=document.id,
                embedding=document.embedding,
                content=document.content,
                metadata=metadata
            )
            
        except Exception as e:
            self.logger.error(f"Error adding document to vector store: {str(e)}")
            raise
    
    async def _save_document_metadata(self, document: Document):
        """Save document metadata to file."""
        try:
            metadata_dir = Path(settings.PROCESSED_DATA_DIR) / "metadata"
            metadata_dir.mkdir(parents=True, exist_ok=True)
            
            metadata_file = metadata_dir / f"{document.id}.json"
            
            # Convert document to dictionary for JSON serialization
            doc_dict = {
                'id': document.id,
                'filename': document.filename,
                'file_path': document.file_path,
                'document_type': document.document_type.value,
                'access_level': document.access_level.value,
                'content_length': len(document.content),
                'metadata': {
                    'title': document.metadata.title,
                    'author': document.metadata.author,
                    'creation_date': document.metadata.creation_date.isoformat() if document.metadata.creation_date else None,
                    'language': document.metadata.language,
                    'subject': document.metadata.subject,
                    'description': document.metadata.description,
                    'keywords': document.metadata.keywords,
                    'file_size': document.metadata.file_size,
                    'page_count': document.metadata.page_count
                },
                'entities': [{
                    'text': e.text,
                    'label': e.label,
                    'entity_type': e.entity_type.value,
                    'start_pos': e.start_pos,
                    'end_pos': e.end_pos,
                    'confidence': e.confidence,
                    'metadata': e.metadata
                } for e in document.entities],
                'created_at': document.created_at.isoformat(),
                'updated_at': document.updated_at.isoformat()
            }
            
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(doc_dict, f, indent=2, ensure_ascii=False)
                
            self.logger.debug(f"Saved metadata for document: {document.id}")
            
        except Exception as e:
            self.logger.error(f"Error saving document metadata: {str(e)}")
    
    async def process_batch(self, directory_path: str, file_pattern: str = "*") -> List[Document]:
        """
        Process multiple documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            file_pattern: File pattern to match (e.g., "*.pdf")
            
        Returns:
            List of processed documents
        """
        processed_documents = []
        
        try:
            directory = Path(directory_path)
            if not directory.exists():
                self.logger.error(f"Directory not found: {directory}")
                return processed_documents
            
            # Find matching files
            files = list(directory.glob(file_pattern))
            self.logger.info(f"Found {len(files)} files to process in {directory}")
            
            for file_path in files:
                if file_path.is_file():
                    document = await self.process_document(str(file_path))
                    if document:
                        processed_documents.append(document)
                    else:
                        self.logger.warning(f"Failed to process: {file_path}")
            
            self.logger.info(f"Successfully processed {len(processed_documents)} documents")
            return processed_documents
            
        except Exception as e:
            self.logger.error(f"Error processing batch: {str(e)}")
            return processed_documents
    
    async def get_document_by_id(self, document_id: str) -> Optional[Document]:
        """
        Retrieve a document by its ID.
        
        Args:
            document_id: Document ID
            
        Returns:
            Document object or None if not found
        """
        try:
            metadata_file = Path(settings.PROCESSED_DATA_DIR) / "metadata" / f"{document_id}.json"
            
            if not metadata_file.exists():
                self.logger.warning(f"Document metadata not found: {document_id}")
                return None
            
            with open(metadata_file, 'r', encoding='utf-8') as f:
                doc_dict = json.load(f)
            
            # Reconstruct document object
            metadata = DocumentMetadata(
                title=doc_dict['metadata'].get('title'),
                author=doc_dict['metadata'].get('author'),
                creation_date=datetime.fromisoformat(doc_dict['metadata']['creation_date']) if doc_dict['metadata'].get('creation_date') else None,
                language=doc_dict['metadata'].get('language', 'en'),
                subject=doc_dict['metadata'].get('subject'),
                description=doc_dict['metadata'].get('description'),
                keywords=doc_dict['metadata'].get('keywords', []),
                file_size=doc_dict['metadata'].get('file_size'),
                page_count=doc_dict['metadata'].get('page_count')
            )
            
            # Reconstruct entities
            entities = []
            for entity_dict in doc_dict.get('entities', []):
                entity = Entity(
                    text=entity_dict['text'],
                    label=entity_dict['label'],
                    entity_type=EntityType(entity_dict['entity_type']),
                    start_pos=entity_dict['start_pos'],
                    end_pos=entity_dict['end_pos'],
                    confidence=entity_dict['confidence'],
                    metadata=entity_dict.get('metadata', {})
                )
                entities.append(entity)
            
            # Read original content if needed (for small files)
            content = ""
            if Path(doc_dict['file_path']).exists():
                doc_type = DocumentType(doc_dict['document_type'])
                content = self._extract_text(Path(doc_dict['file_path']), doc_type)
            
            document = Document(
                id=doc_dict['id'],
                filename=doc_dict['filename'],
                file_path=doc_dict['file_path'],
                document_type=DocumentType(doc_dict['document_type']),
                access_level=AccessLevel(doc_dict['access_level']),
                content=content,
                metadata=metadata,
                entities=entities,
                created_at=datetime.fromisoformat(doc_dict['created_at']),
                updated_at=datetime.fromisoformat(doc_dict['updated_at'])
            )
            
            return document
            
        except Exception as e:
            self.logger.error(f"Error retrieving document {document_id}: {str(e)}")
            return None
    
    async def delete_document(self, document_id: str) -> bool:
        """
        Delete a document from the index.
        
        Args:
            document_id: Document ID to delete
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Delete from vector store
            await self.vector_store.delete_document(document_id)
            
            # Delete metadata file
            metadata_file = Path(settings.PROCESSED_DATA_DIR) / "metadata" / f"{document_id}.json"
            if metadata_file.exists():
                metadata_file.unlink()
            
            self.logger.info(f"Deleted document: {document_id}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error deleting document {document_id}: {str(e)}")
            return False
    
    async def get_processing_stats(self) -> Dict[str, Any]:
        """
        Get statistics about processed documents.
        
        Returns:
            Dictionary with processing statistics
        """
        try:
            metadata_dir = Path(settings.PROCESSED_DATA_DIR) / "metadata"
            
            if not metadata_dir.exists():
                return {"total_documents": 0}
            
            # Count metadata files
            metadata_files = list(metadata_dir.glob("*.json"))
            total_docs = len(metadata_files)
            
            # Analyze document types and other stats
            doc_types = {}
            access_levels = {}
            languages = {}
            total_entities = 0
            
            for metadata_file in metadata_files:
                try:
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        doc_dict = json.load(f)
                    
                    # Count by document type
                    doc_type = doc_dict.get('document_type', 'unknown')
                    doc_types[doc_type] = doc_types.get(doc_type, 0) + 1
                    
                    # Count by access level
                    access_level = doc_dict.get('access_level', 'unknown')
                    access_levels[access_level] = access_levels.get(access_level, 0) + 1
                    
                    # Count by language
                    language = doc_dict.get('metadata', {}).get('language', 'unknown')
                    languages[language] = languages.get(language, 0) + 1
                    
                    # Count entities
                    total_entities += len(doc_dict.get('entities', []))
                    
                except Exception as e:
                    self.logger.warning(f"Error reading metadata file {metadata_file}: {str(e)}")
                    continue
            
            # Get vector store stats
            vector_stats = await self.vector_store.get_collection_stats()
            
            return {
                "total_documents": total_docs,
                "document_types": doc_types,
                "access_levels": access_levels,
                "languages": languages,
                "total_entities": total_entities,
                "average_entities_per_doc": total_entities / total_docs if total_docs > 0 else 0,
                "vector_store": vector_stats
            }
            
        except Exception as e:
            self.logger.error(f"Error getting processing stats: {str(e)}")
            return {"error": str(e)}